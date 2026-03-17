import io
import fitz
from docx import Document


def read_pdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(stream=content, filetype="pdf")
    parts = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n".join(parts).strip()


def read_docx(content: bytes) -> str:
    """Extract text from Word document using python-docx."""
    doc = Document(io.BytesIO(content))
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def read_txt(content: bytes) -> str:
    """Read plain text file."""
    return content.decode("utf-8", errors="replace").strip()
